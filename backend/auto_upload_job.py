#!/usr/bin/env python
"""
自动职位上传程序
处理docx文件的AI解析和数据库存储
"""

import os
import sys
import logging
import json
import requests
from datetime import datetime
from pathlib import Path
import django
from django.conf import settings

# 添加Django项目路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 设置Django环境
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from jobs.models import Job
from matching.google_ai_service import GoogleAIService
from error_handler import create_error_handler, DEFAULT_ERROR_CONFIG
from users.models import User, EmployerProfile  # 导入EmployerProfile
import docx
import PyPDF2
from io import BytesIO
import traceback


class AutoJobUploader:
    """自动职位上传处理器"""
    
    def __init__(self, log_level=logging.INFO):
        """初始化自动上传器"""
        self.setup_logging(log_level)
        
        # 初始化错误处理器
        error_config = DEFAULT_ERROR_CONFIG.copy()
        error_config['notification']['email_recipients'] = ['admin@company.com']  # 可配置
        self.error_handler = create_error_handler(error_config)
        
        self.ai_service = GoogleAIService()
        self.logger = logging.getLogger(__name__)
        
    def setup_logging(self, log_level):
        """设置日志系统"""
        # 创建日志目录
        log_dir = Path('logs')
        log_dir.mkdir(exist_ok=True)
        
        # 设置日志文件路径
        log_file = log_dir / f'auto_upload_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'
        
        # 配置日志格式
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        
        # 配置文件处理器
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(log_level)
        file_handler.setFormatter(formatter)
        
        # 配置控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setLevel(log_level)
        console_handler.setFormatter(formatter)
        
        # 配置根日志器
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(log_level)
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
        
        self.logger.info(f"日志系统初始化完成，日志文件: {log_file}")
        return log_file
        
    def extract_text_from_docx(self, file_path):
        """从docx文件提取文本内容"""
        try:
            self.logger.info(f"开始提取docx文件内容: {file_path}")
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_content)
            self.logger.info(f"成功提取文本内容，长度: {len(full_text)} 字符")
            self.logger.debug(f"提取的文本内容预览: {full_text[:200]}...")
            
            return full_text
            
        except Exception as e:
            self.logger.error(f"提取docx文件内容失败: {str(e)}")
            self.logger.error(f"错误详情: {traceback.format_exc()}")
            raise
    
    def parse_job_with_ai(self, text_content, filename):
        """使用AI解析职位描述"""
        context = {
            'operation': 'ai_parsing',
            'filename': filename,
            'content_length': len(text_content)
        }
        
        try:
            # 使用错误处理器执行AI解析
            result = self.error_handler.with_error_handling(
                lambda: self._parse_job_with_ai_internal(text_content, filename),
                context=context,
                retry=True,
                circuit_breaker=True,  # 使用熔断器防止AI服务过载
                notify_on_error=False  # AI解析失败不需要立即通知
            )
            
            return result
            
        except Exception as e:
            self.logger.error(f"AI解析失败: {filename} - {str(e)}")
            
            # 记录错误但不通知（AI解析失败是可恢复的）
            self.error_handler.handle_error(
                e, 
                context=context, 
                notify=False, 
                attempt_recovery=False
            )
            
            # 返回默认数据结构
            return {
                'success': False,
                'data': self._get_default_job_data(),
                'error': str(e)
            }
    
    def _parse_job_with_ai_internal(self, text_content, filename):
        """内部AI解析逻辑"""
        self.logger.info(f"开始AI解析职位描述: {filename}")
        
        if not self.ai_service.is_enabled:
            self.logger.error("AI服务未启用，无法进行解析")
            raise Exception("AI服务未配置或不可用")
        
        # 构建AI提示词
        prompt = self._build_job_parsing_prompt(text_content)
        self.logger.debug(f"使用的AI提示词: {prompt[:300]}...")
        
        # 调用AI服务
        ai_response = self.ai_service.parse_job_description(text_content)
        self.logger.info("AI解析完成，开始解析响应")
        self.logger.debug(f"AI原始响应: {ai_response}")
        
        if not ai_response:
            raise ValueError("AI服务返回空响应")
        
        # 解析AI响应
        parsed_data = self._parse_ai_response(ai_response)
        self.logger.info("AI响应解析完成")
        self.logger.info(f"解析结果: {json.dumps(parsed_data, ensure_ascii=False, indent=2)}")
        
        return {
            'success': True,
            'data': parsed_data,
            'ai_response': ai_response,
            'prompt_used': prompt
        }
    
    def _build_job_parsing_prompt(self, text_content):
        """构建职位解析AI提示词"""
        return f"""
You are an expert job description parser. Please carefully analyze the following job posting content and extract comprehensive information.

Job Posting Content:
{text_content}

Please extract the following information systematically. For each field, provide the most accurate information found. If information is not available, use appropriate defaults:

Job Title: [Extract the job title/position name]
Company: [Extract company name if mentioned]
Location: [Extract job location/city]
Job Type: [Extract employment type - use only: full_time, part_time, contract, internship, freelance]
Experience Level: [Extract required experience level - use only: entry, junior, mid, senior, lead, executive]
Remote Option: [Extract work arrangement - use only: remote, hybrid, on_site]
Salary Min: [Extract minimum salary if mentioned, otherwise leave empty]
Salary Max: [Extract maximum salary if mentioned, otherwise leave empty]
Description: [Extract detailed job description and responsibilities]
Requirements: [Extract required skills, qualifications, and experience]
Responsibilities: [Extract job responsibilities and duties]
Benefits: [Extract benefits, perks, and compensation details - separate multiple items with commas]
Application Deadline: [Extract application deadline if mentioned, format as YYYY-MM-DD]

IMPORTANT INSTRUCTIONS:
1. Be thorough and accurate in extraction
2. For job_type, use only: full_time, part_time, contract, internship, freelance
3. For experience_level, use only: entry, junior, mid, senior, lead, executive
4. For remote_option, use only: remote, hybrid, on_site
5. Keep original formatting and details where possible
6. If salary is mentioned as a range, extract both min and max
7. For dates, use YYYY-MM-DD format
8. For benefits, separate multiple items with commas

Please format your response exactly as shown above with clear labels for each field.
"""
    
    def _parse_ai_response(self, ai_response):
        """解析AI响应为结构化数据"""
        parsed_data = {
            'title': '',
            'company': '',
            'location_city': 'Not Specified',
            'location_state': '',
            'location_country': '中国',
            'job_type': 'full_time',
            'experience_level': 'entry',
            'remote_option': 'on_site',
            'salary_min': None,
            'salary_max': None,
            'description': '',
            'requirements': '',
            'responsibilities': '',  # 添加职责字段
            'benefits': [],  # 改为列表格式以匹配JSONField
            'application_deadline': None,  # 改为None以便后续处理
            'category_id': 1,  # 默认分类ID，需要确保数据库中存在
        }
        
        try:
            lines = ai_response.strip().split('\n')
            for line in lines:
                line = line.strip()
                if ':' in line:
                    key, value = line.split(':', 1)
                    key = key.strip().lower()
                    value = value.strip()
                    
                    if 'job title' in key:
                        parsed_data['title'] = value
                    elif 'company' in key:
                        parsed_data['company'] = value
                    elif 'location' in key:
                        parsed_data['location_city'] = value if value != 'Not Specified' else 'Not Specified'
                    elif 'job type' in key:
                        if value.lower() in ['full_time', 'part_time', 'contract', 'internship']:
                            parsed_data['job_type'] = value.lower()
                    elif 'experience level' in key:
                        if value.lower() in ['entry', 'mid', 'senior', 'executive']:
                            parsed_data['experience_level'] = value.lower()
                    elif 'remote option' in key:
                        if value.lower() in ['remote', 'hybrid', 'on_site']:
                            parsed_data['remote_option'] = value.lower()
                    elif 'salary min' in key:
                        # 处理薪资字段，转换为数字或None
                        if value and value.lower() not in ['not provided', 'not specified', '']:
                            import re
                            numbers = re.findall(r'\d+', value)
                            if numbers:
                                try:
                                    parsed_data['salary_min'] = int(numbers[0])
                                except (ValueError, IndexError):
                                    parsed_data['salary_min'] = None
                            else:
                                parsed_data['salary_min'] = None
                        else:
                            parsed_data['salary_min'] = None
                    elif 'salary max' in key:
                        # 处理薪资字段，转换为数字或None
                        if value and value.lower() not in ['not provided', 'not specified', '']:
                            import re
                            numbers = re.findall(r'\d+', value)
                            if numbers:
                                try:
                                    parsed_data['salary_max'] = int(numbers[0])
                                except (ValueError, IndexError):
                                    parsed_data['salary_max'] = None
                            else:
                                parsed_data['salary_max'] = None
                        else:
                            parsed_data['salary_max'] = None
                    elif 'description' in key:
                        parsed_data['description'] = value
                    elif 'requirements' in key:
                        parsed_data['requirements'] = value
                    elif 'responsibilities' in key:
                        parsed_data['responsibilities'] = value
                    elif 'benefits' in key:
                        # 将benefits转换为列表格式
                        if value and value.lower() not in ['not provided', 'not specified', '']:
                            # 简单的分割处理，可以根据需要改进
                            benefits_list = [item.strip() for item in value.split(',') if item.strip()]
                            parsed_data['benefits'] = benefits_list
                        else:
                            parsed_data['benefits'] = []
                    elif 'application deadline' in key:
                        # 处理申请截止日期
                        if value and value.lower() not in ['not provided', 'not specified', '']:
                            try:
                                from datetime import datetime
                                # 尝试解析日期格式
                                if '-' in value:
                                    parsed_date = datetime.strptime(value, '%Y-%m-%d')
                                    parsed_data['application_deadline'] = parsed_date
                                else:
                                    parsed_data['application_deadline'] = None
                            except ValueError:
                                parsed_data['application_deadline'] = None
                        else:
                            parsed_data['application_deadline'] = None
            
            return parsed_data
            
        except Exception as e:
            self.logger.error(f"解析AI响应失败: {str(e)}")
            return parsed_data
    
    def _get_default_job_data(self):
        """获取默认职位数据"""
        return {
            'title': 'AI Product Engineer',
            'company': 'Tech Company',
            'location_city': 'Not Specified',
            'location_state': '',
            'location_country': '中国',
            'job_type': 'full_time',
            'experience_level': 'entry',
            'remote_option': 'on_site',
            'salary_min': None,
            'salary_max': None,
            'description': 'Job description not available',
            'requirements': 'Requirements not specified',
            'responsibilities': 'Responsibilities not specified',
            'benefits': [],
            'application_deadline': None,
            'category_id': 1,
        }
    
    def save_to_database(self, job_data, filename):
        """保存职位数据到数据库"""
        context = {
            'operation': 'database_save',
            'filename': filename,
            'job_title': job_data.get('title', 'Unknown')
        }
        
        try:
            # 使用错误处理器执行数据库保存
            result = self.error_handler.with_error_handling(
                lambda: self._save_to_database_internal(job_data, filename),
                context=context,
                retry=True,
                notify_on_error=True  # 数据库错误需要通知
            )
            
            return result
            
        except Exception as e:
            self.logger.error(f"数据库保存失败: {filename} - {str(e)}")
            
            # 记录错误并通知（数据库错误是严重的）
            self.error_handler.handle_error(
                e, 
                context=context, 
                notify=True, 
                attempt_recovery=True
            )
            
            raise  # 重新抛出异常，因为数据库保存失败是不可恢复的
    
    def _save_to_database_internal(self, job_data, filename):
        """内部数据库保存逻辑"""
        self.logger.info(f"开始保存职位数据到数据库: {filename}")
        
        # 获取或创建默认用户（雇主）
        user, created = User.objects.get_or_create(
            username='auto_uploader',
            defaults={
                'email': 'auto@example.com',
                'first_name': 'Auto',
                'last_name': 'Uploader',
                'user_type': 'employer'  # 设置为雇主类型
            }
        )
        
        if created:
            self.logger.info("创建了新的自动上传用户")
        
        # 获取或创建雇主档案
        from users.models import EmployerProfile
        employer_profile, profile_created = EmployerProfile.objects.get_or_create(
            user=user,
            defaults={
                'company_name': job_data.get('company', 'Auto Upload Company'),
                'company_description': '自动上传系统创建的公司',
                'industry': 'technology',
                'company_size': 'medium',
                'contact_person': 'Auto Uploader',
                'contact_title': 'System Admin',
                'office_address': '自动生成地址'
            }
        )
        
        # 验证必需字段
        required_fields = ['title']
        for field in required_fields:
            if not job_data.get(field):
                raise ValueError(f"缺少必需字段: {field}")
        
        # 获取或创建职位分类
        from jobs.models import JobCategory
        try:
            category = JobCategory.objects.get(id=job_data.get('category_id', 1))
        except JobCategory.DoesNotExist:
            # 如果指定的分类不存在，创建一个默认分类
            category, _ = JobCategory.objects.get_or_create(
                name='技术类',
                defaults={
                    'description': '技术相关职位',
                    'is_active': True
                }
            )
        
        # 创建职位记录
        job = Job.objects.create(
            title=job_data['title'],
            description=job_data['description'],
            requirements=job_data['requirements'],
            responsibilities=job_data.get('responsibilities', ''),
            job_type=job_data['job_type'],
            experience_level=job_data['experience_level'],
            location_city=job_data['location_city'],
            location_state=job_data.get('location_state', ''),
            location_country=job_data.get('location_country', '中国'),
            remote_option=job_data['remote_option'],
            salary_min=job_data['salary_min'] if job_data['salary_min'] else None,
            salary_max=job_data['salary_max'] if job_data['salary_max'] else None,
            benefits=job_data.get('benefits', []),
            application_deadline=job_data['application_deadline'] if job_data['application_deadline'] else None,
            employer=employer_profile,
            category=category,
            is_active=True
        )
        
        self.logger.info(f"成功保存职位到数据库，ID: {job.id}")
        self.logger.info(f"职位详情: {job.title} - {job.employer.company_name} - {job.location_city}")
        
        return job
    
    def process_file(self, file_path):
        """处理单个文件的完整流程"""
        file_path = Path(file_path)
        self.logger.info(f"开始处理文件: {file_path}")
        
        context = {
            'file_path': str(file_path),
            'operation': 'file_processing'
        }
        
        try:
            # 使用错误处理器执行文件处理
            result = self.error_handler.with_error_handling(
                lambda: self._process_file_internal(file_path),
                context=context,
                retry=True,
                notify_on_error=True
            )
            
            self.logger.info(f"文件处理成功: {file_path}")
            return result
            
        except Exception as e:
            self.logger.error(f"文件处理失败: {file_path} - {str(e)}")
            
            # 记录详细错误信息
            self.error_handler.handle_error(
                e, 
                context=context, 
                notify=True, 
                attempt_recovery=False
            )
            
            error_result = {
                'success': False,
                'file_path': str(file_path),
                'error': str(e),
                'error_type': type(e).__name__,
                'processing_time': datetime.now().isoformat()
            }
            
            return error_result
    
    def _process_file_internal(self, file_path):
        """内部文件处理逻辑"""
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if not file_path.suffix.lower() == '.docx':
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
        
        # 步骤1: 提取文本内容
        self.logger.info("步骤1: 提取文档内容")
        text_content = self.extract_text_from_docx(file_path)
        
        # 步骤2: AI解析
        self.logger.info("步骤2: AI解析职位描述")
        ai_result = self.parse_job_with_ai(text_content, file_path.name)
        
        if not ai_result['success']:
            self.logger.warning("AI解析失败，使用默认数据")
        
        # 步骤3: 保存到数据库
        self.logger.info("步骤3: 保存到数据库")
        job = self.save_to_database(ai_result['data'], file_path.name)
        
        # 记录成功结果
        result = {
            'success': True,
            'file_path': str(file_path),
            'job_id': job.id,
            'job_title': job.title,
            'ai_enabled': ai_result['success'],
            'processing_time': datetime.now().isoformat()
        }
        
        self.logger.info("文件处理完成")
        self.logger.info(f"处理结果: {json.dumps(result, ensure_ascii=False, indent=2)}")
        
        return result


def main():
    """主函数"""
    if len(sys.argv) != 2:
        print("使用方法: python auto_upload_job.py <docx文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # 创建自动上传器
    uploader = AutoJobUploader(log_level=logging.INFO)
    
    # 处理文件
    result = uploader.process_file(file_path)
    
    # 输出最终结果
    print("\n" + "="*50)
    print("处理结果:")
    print("="*50)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    
    if result['success']:
        print(f"\n✅ 成功处理文件: {file_path}")
        print(f"📝 创建职位: {result['job_title']} (ID: {result['job_id']})")
        print(f"🤖 AI解析: {'成功' if result['ai_enabled'] else '失败（使用默认数据）'}")
    else:
        print(f"\n❌ 处理失败: {file_path}")
        print(f"🚫 错误: {result['error']}")
        sys.exit(1)


if __name__ == '__main__':
    main()